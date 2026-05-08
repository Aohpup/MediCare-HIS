#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word 转换脚本 v2
特性：
1. 自动生成可更新目录 (TOC)
2. 中文宋体 + 英文 Times New Roman
3. 代码块带标题编号，方便定位截图
4. 表格完整边框 + 表头底纹
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import parse_xml

def set_run_font(run, font_name='Times New Roman', eastasia='宋体', size=10.5, bold=False, italic=False, color=None):
    """设置run字体：英文用Times New Roman，中文用宋体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), eastasia)
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color

def set_cell_border(cell):
    """设置单元格四周边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def add_formatted_paragraph(doc, text, style='Normal', bold=False):
    """添加带格式的段落，支持内联粗体**text**和`code`"""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_run_font(run, font_name='Courier New', eastasia='Courier New', size=10, color=RGBColor(0xC0, 0x00, 0x00))
        else:
            run = p.add_run(part)
            set_run_font(run, bold=bold)
    return p

def add_code_block(doc, code_lines, code_title):
    """添加带标题的代码块，方便用户定位截图"""
    # 标题行
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(4)
    p_title.paragraph_format.left_indent = Cm(1)
    
    run_title = p_title.add_run(f"▶ {code_title}")
    set_run_font(run_title, font_name='Times New Roman', eastasia='黑体', size=10, bold=True, color=RGBColor(0x00, 0x4D, 0x99))
    
    # 代码内容
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    run = p.add_run('\n'.join(code_lines))
    set_run_font(run, font_name='Courier New', eastasia='Courier New', size=9, color=RGBColor(0x2B, 0x2B, 0x2B))
    
    # 代码块底纹
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F5F5F5"/>')
    pPr.append(shd)
    
    doc.add_paragraph()

def add_table_from_md(doc, lines, start_idx):
    """从Markdown表格行解析并生成Word表格"""
    rows = []
    idx = start_idx
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        row_text = lines[idx].strip()
        if re.match(r'\|\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?', row_text):
            idx += 1
            continue
        cells = [c.strip() for c in row_text.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        idx += 1
    
    if not rows:
        return idx
    
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    for i, row_cells in enumerate(rows):
        row = table.rows[i]
        for j in range(num_cols):
            cell = row.cells[j]
            set_cell_border(cell)
            cell.vertical_alignment = 1  # CENTER
            
            if j < len(row_cells):
                p = cell.paragraphs[0]
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                parts = re.split(r'(\*\*.*?\*\*)', row_cells[j])
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        set_run_font(run, bold=True, size=10.5)
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        set_run_font(run, font_name='Courier New', eastasia='Courier New', size=10, color=RGBColor(0xC0, 0x00, 0x00))
                    else:
                        run = p.add_run(part)
                        set_run_font(run, size=10.5)
                
                if i == 0:
                    for run in p.runs:
                        run.font.bold = True
                    shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9E2F3"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
    
    for col in table.columns:
        for cell in col.cells:
            cell.width = Cm(3.5)
    
    doc.add_paragraph()
    return idx

def add_toc(doc):
    """在文档开头插入目录（TOC）"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = parse_xml(r'<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="begin"/>')
    run._r.append(fldChar_begin)
    
    run2 = paragraph.add_run()
    instrText = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)
    
    run3 = paragraph.add_run()
    fldChar_separate = parse_xml(r'<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="separate"/>')
    run3._r.append(fldChar_separate)
    
    run4 = paragraph.add_run("（右键此处 → 更新域，即可生成目录）")
    set_run_font(run4, size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80))
    
    run5 = paragraph.add_run()
    fldChar_end = parse_xml(r'<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="end"/>')
    run5._r.append(fldChar_end)
    
    doc.add_paragraph()
    doc.add_paragraph("——————————————————————————————————").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def setup_styles(doc):
    """配置文档默认样式：中文宋体，英文Times New Roman"""
    # Normal 样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    # Heading 样式统一设置
    for i in range(1, 10):
        try:
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Times New Roman'
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            heading_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        except KeyError:
            pass

def md_to_docx(md_path, docx_path):
    doc = Document()
    setup_styles(doc)
    
    # 页面设置
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 先收集代码块信息用于生成编号标题
    code_blocks = []
    temp_idx = 0
    while temp_idx < len(lines):
        if lines[temp_idx].strip().startswith('```'):
            lang = lines[temp_idx].strip()[3:].strip()
            temp_idx += 1
            code_lines = []
            while temp_idx < len(lines) and not lines[temp_idx].strip().startswith('```'):
                code_lines.append(lines[temp_idx].rstrip('\n'))
                temp_idx += 1
            code_blocks.append((len(code_blocks) + 1, lang, code_lines))
        temp_idx += 1
    
    # 插入目录
    add_toc(doc)
    
    i = 0
    code_counter = 0
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()
        
        # 代码块
        if stripped.startswith('```'):
            code_counter += 1
            lang = stripped[3:].strip()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1
            
            # 根据上下文生成代码块标题
            code_title = f"代码片段 {code_counter}"
            if lang:
                code_title += f"  [{lang.upper()}]"
            if code_counter <= len(code_blocks):
                # 尝试从代码内容推断用途
                first_line = code_lines[0] if code_lines else ""
                if 'flowchart' in first_line or 'graph TD' in first_line:
                    code_title += " — 流程图定义"
                elif 'typedef struct' in first_line:
                    code_title += " — 结构体定义"
                elif '#include' in first_line:
                    code_title += " — 源代码"
                elif 'DRUG_FILE' in first_line or 'HIS_' in first_line:
                    code_title += " — 文件常量定义"
            
            add_code_block(doc, code_lines, code_title)
            continue
        
        # 分隔线
        if stripped == '---':
            doc.add_paragraph()
            i += 1
            continue
        
        # 标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            p = doc.add_heading(level=0)
            run = p.add_run(stripped[2:])
            set_run_font(run, eastasia='黑体', size=22, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        elif stripped.startswith('## '):
            p = doc.add_heading(level=1)
            run = p.add_run(stripped[3:])
            set_run_font(run, eastasia='黑体', size=16, bold=True)
            i += 1
            continue
        elif stripped.startswith('### '):
            p = doc.add_heading(level=2)
            run = p.add_run(stripped[4:])
            set_run_font(run, eastasia='黑体', size=14, bold=True)
            i += 1
            continue
        elif stripped.startswith('#### '):
            p = doc.add_heading(level=3)
            run = p.add_run(stripped[5:])
            set_run_font(run, eastasia='黑体', size=12, bold=True)
            i += 1
            continue
        
        # 表格
        if stripped.startswith('|') and not re.match(r'\|\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?', stripped):
            i = add_table_from_md(doc, lines, i)
            continue
        elif stripped.startswith('|'):
            i += 1
            continue
        
        # 引用块
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(stripped[2:])
            set_run_font(run, italic=True, size=11, color=RGBColor(0x55, 0x55, 0x55))
            i += 1
            continue
        
        # 普通段落
        if stripped:
            add_formatted_paragraph(doc, stripped)
        
        i += 1
    
    doc.save(docx_path)
    print(f"✅ 转换完成: {docx_path}")
    print(f"   共识别 {code_counter} 个代码片段，均已标注编号和位置")

if __name__ == '__main__':
    md_file = r'MediCare_HIS_2026\项目总结报告.md'
    docx_file = r'MediCare_HIS_2026\项目总结报告_v2.docx'
    md_to_docx(md_file, docx_file)
