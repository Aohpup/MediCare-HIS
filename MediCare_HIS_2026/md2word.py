#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word 转换脚本
特点：表格带完整边框、表头加粗、自动列宽、中文友好
"""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import parse_xml

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                          r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          r'</w:tcBorders>')
    tcPr.append(tcBorders)

def set_run_font(run, font_name='宋体', size=10.5, bold=False, italic=False, color=None):
    """设置run字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False):
    """添加带格式的段落，支持内联粗体**text**"""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    
    # 解析内联粗体 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_run_font(run, font_name='Courier New', size=10)
        else:
            run = p.add_run(part)
            set_run_font(run, bold=bold)
    return p

def add_table_from_md(doc, lines, start_idx):
    """从Markdown表格行解析并生成Word表格"""
    rows = []
    idx = start_idx
    # 收集表格所有行
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        row_text = lines[idx].strip()
        # 跳过分隔行 |---|---|
        if re.match(r'\|\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?', row_text):
            idx += 1
            continue
        cells = [c.strip() for c in row_text.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        idx += 1
    
    if not rows:
        return idx
    
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    # 设置表格整体宽度为页面可用宽度
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    
    for i, row_cells in enumerate(rows):
        row = table.rows[i]
        for j in range(num_cols):
            cell = row.cells[j]
            set_cell_border(cell)
            cell.vertical_alignment = 1  # CENTER
            
            if j < len(row_cells):
                cell_text = row_cells[j]
                # 清除默认段落
                p = cell.paragraphs[0]
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 解析单元格内的粗体
                parts = re.split(r'(\*\*.*?\*\*)', cell_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        set_run_font(run, bold=True, size=10.5)
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        set_run_font(run, font_name='Courier New', size=10)
                    else:
                        run = p.add_run(part)
                        set_run_font(run, size=10.5)
                
                # 第一行作为表头加粗并设置底纹
                if i == 0:
                    for run in p.runs:
                        run.font.bold = True
                    # 设置表头底纹
                    shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9E2F3"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
    
    # 尝试平均分配列宽
    for col in table.columns:
        for cell in col.cells:
            cell.width = Cm(3.5)
    
    doc.add_paragraph()  # 表格后空一行
    return idx

def add_code_block(doc, lines, start_idx):
    """添加代码块"""
    idx = start_idx + 1
    code_lines = []
    while idx < len(lines) and not lines[idx].strip().startswith('```'):
        code_lines.append(lines[idx])
        idx += 1
    idx += 1  # 跳过结束 ```
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    run = p.add_run('\n'.join(code_lines))
    set_run_font(run, font_name='Courier New', size=9, color=RGBColor(0x2B, 0x2B, 0x2B))
    
    # 代码块底纹
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F5F5F5"/>')
    pPr.append(shd)
    
    doc.add_paragraph()
    return idx

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    # 页面设置：A4，适当边距
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()
        
        # 代码块
        if stripped.startswith('```'):
            i = add_code_block(doc, lines, i)
            continue
        
        # 分隔线
        if stripped == '---':
            doc.add_paragraph().paragraph_format.border_bottom = True
            i += 1
            continue
        
        # 标题
        if stripped.startswith('# '):
            p = doc.add_heading(level=0)
            run = p.add_run(stripped[2:])
            set_run_font(run, font_name='黑体', size=22, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        elif stripped.startswith('## '):
            p = doc.add_heading(level=1)
            run = p.add_run(stripped[3:])
            set_run_font(run, font_name='黑体', size=16, bold=True)
            i += 1
            continue
        elif stripped.startswith('### '):
            p = doc.add_heading(level=2)
            run = p.add_run(stripped[4:])
            set_run_font(run, font_name='黑体', size=14, bold=True)
            i += 1
            continue
        elif stripped.startswith('#### '):
            p = doc.add_heading(level=3)
            run = p.add_run(stripped[5:])
            set_run_font(run, font_name='黑体', size=12, bold=True)
            i += 1
            continue
        
        # 表格
        if stripped.startswith('|') and not re.match(r'\|\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?', stripped):
            i = add_table_from_md(doc, lines, i)
            continue
        elif stripped.startswith('|'):
            # 单独遇到分隔行，跳过
            i += 1
            continue
        
        # 引用块
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(stripped[2:])
            set_run_font(run, italic=True, size=11, color=RGBColor(0x55, 0x55, 0x55))
            i += 1
            continue
        
        # 普通段落
        if stripped:
            add_formatted_paragraph(doc, stripped)
        else:
            # 空行用于分段
            pass
        
        i += 1
    
    doc.save(docx_path)
    print(f"✅ 转换完成: {docx_path}")

if __name__ == '__main__':
    md_file = r'MediCare_HIS_2026\项目总结报告.md'
    docx_file = r'MediCare_HIS_2026\项目总结报告.docx'
    md_to_docx(md_file, docx_file)
